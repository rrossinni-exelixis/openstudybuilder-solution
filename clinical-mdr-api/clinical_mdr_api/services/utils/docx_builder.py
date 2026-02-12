import io
import logging
import os
from functools import reduce
from itertools import zip_longest
from typing import Any, Mapping

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.blkcntnr import BlockItemContainer
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.section import Section
from docx.shared import Inches
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph

from common.telemetry import trace_calls

BLOCK_ELEMENTS_TO_PARAGRAPH = {"p", "div", "li", "dd", "dt"}
INLINE_ELEMENTS_TO_FONT_PROPERTIES = {"b", "em", "i", "sub", "sup", "s", "strong", "u"}

log = logging.getLogger(__name__)


class DocxBuilder:
    DEFAULT_TEMPLATE_FILENAME = "template.docx"

    def __init__(
        self,
        styles: Mapping | None = None,
        landscape: bool = False,
        margins: list[float] | None = None,
        template: str | None = None,
    ):
        template_filename = os.path.join(
            os.path.dirname(__file__),
            template or self.DEFAULT_TEMPLATE_FILENAME,
        )
        self.document = self.load_document(template_filename)
        self.styles: dict[Any, Any] = {}
        self.clear_document()
        if styles:
            self.create_styles(styles)
        if landscape:
            self.set_landscape()
        if margins:
            assert len(margins) == 4, "margins should be a sequence of 4 floats"
            self.set_margins([Inches(i) for i in margins])

    def set_margins(self, margins) -> Section:
        section = self.document.sections[-1]
        section.top_margin = min(section.top_margin, margins[0])
        section.right_margin = min(section.right_margin, margins[1])
        section.bottom_margin = min(section.bottom_margin, margins[2])
        section.left_margin = min(section.left_margin, margins[3])
        return section

    def set_landscape(self):
        # Set page layout to Landscape
        # pylint: disable=no-member
        section = self.document.sections[-1]
        if section.orientation != WD_ORIENTATION.LANDSCAPE:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            page_width = section.page_width
            section.page_width = section.page_height
            section.page_height = page_width
        return section

    def create_styles(self, styles: Mapping):
        # Ensure styles are defined
        for key, value in styles.items():
            name, typ = value
            if key not in self.styles:
                self.styles[key] = value
            if name not in self.document.styles:
                log.debug("Creating %s style: %s", typ, name)
                self.document.styles.add_style(name, typ)

    def clear_document(self):
        # pylint: disable=protected-access
        log.debug("Removing all content from document")
        self.document._element.body.clear_content()

    @staticmethod
    @trace_calls(args=[0], kwargs=["template_filename"])
    def load_document(template_filename: str) -> DocumentObject:
        log.debug("Reading document: %s", template_filename)
        with open(template_filename, "rb") as file:
            return Document(file)

    def create_table(self, num_rows: int, num_columns: int) -> Table:
        table = self.document.add_table(rows=num_rows, cols=num_columns)
        table.autofit = True
        style = self.styles.get("table")
        if style:
            table.style = self.document.styles[style[0]]
        return table

    @staticmethod
    def add_row(table: Table, cell_text_content: list[str | None]) -> _Row:
        """Adds a row to the table and fills the cells text content"""
        row = table.add_row()
        for i, text in enumerate(cell_text_content):
            if text is not None:
                row.cells[i].text = text
        return row

    @staticmethod
    def format_row(row: _Row, styles: list[str] | None = None, **kwargs):
        """Applies styles on each cell of a table row, and sets paragraph formattings from kwargs"""
        for cell, style in zip_longest(row.cells, styles or []):
            paragraph = cell.paragraphs[0]
            if style is not None:
                paragraph.style = style
            for k, value in kwargs.items():
                setattr(paragraph, k, value)

    @staticmethod
    def repeat_table_header(row: _Row):
        """Set table row to repeat on every page"""
        # pylint: disable=protected-access
        table_row = row._tr
        tr_pr = table_row.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)

    @staticmethod
    def set_vertical_cell_direction(cell: _Cell, direction: str):
        """Set text direction to vertical for a cell: tbRl -- top to bottom, btLr -- bottom to top"""
        # pylint: disable=protected-access
        assert direction in ("tbRl", "btLr")
        table_cell = cell._tc
        tc_pr = table_cell.get_or_add_tcPr()
        text_direction = OxmlElement("w:textDirection")
        text_direction.set(qn("w:val"), direction)  # btLr tbRl
        tc_pr.append(text_direction)

    @staticmethod
    def merge_cells(cells: tuple[_Cell, ...]) -> _Cell:
        return reduce(lambda c1, c2: c1.merge(c2), cells)

    def get_document_stream(self) -> io.BytesIO:
        doc = io.BytesIO()
        self.document.save(doc)
        doc.seek(0)
        return doc

    @staticmethod
    def delete_paragraph(paragraph: Paragraph) -> None:
        element = paragraph._element  # pylint: disable=protected-access
        element.getparent().remove(element)
        element._p = element._element = None

    def add_html(
        self,
        container: BlockItemContainer,
        markup: str,
        default_style: str | None = None,
    ) -> None:
        style_map = self.styles.copy()
        style_map["p"] = style_map[default_style]

        soup = BeautifulSoup(markup, "lxml")
        if not soup.body:
            # if markup was an empty string or just whitespaces
            return
        # Always has a .body with minimum .contents as a <p> tag, even if markup was just a simple plain-text string

        def recursion(children, paragraph, paragraph_style, run_styles):
            for element in children:
                if isinstance(element, Tag):
                    paragraph_style = style_map.get(element.name, (paragraph_style,))[0]

                    if element.name in BLOCK_ELEMENTS_TO_PARAGRAPH:
                        paragraph = container.add_paragraph(style=paragraph_style)
                        recursion(element.children, paragraph, paragraph_style, set())

                    elif element.name in INLINE_ELEMENTS_TO_FONT_PROPERTIES:
                        recursion(
                            element.children,
                            paragraph,
                            paragraph_style,
                            run_styles | {element.name},
                        )

                    elif element.name == "br":
                        if paragraph:
                            run = paragraph.add_run("break")
                            run.add_break(WD_BREAK.LINE)

                    else:
                        recursion(
                            element.children, paragraph, paragraph_style, run_styles
                        )

                elif isinstance(element, NavigableString):
                    text = str(element) or None
                    if text and paragraph:
                        run = paragraph.add_run(text)
                        run.font.bold = "b" in run_styles
                        run.font.bold = "strong" in run_styles
                        run.font.italic = "i" in run_styles
                        run.font.italic = "em" in run_styles
                        run.font.strike = "s" in run_styles
                        run.font.subscript = "sub" in run_styles
                        run.font.superscript = "sup" in run_styles
                        run.font.underline = "u" in run_styles

        recursion(
            soup.body.children, paragraph=None, paragraph_style=None, run_styles=set()
        )

    def replace_content(
        self, container: BlockItemContainer, text: str, style: str | None = None
    ):
        style = self.styles.get(style, (None,))[0]
        container.add_paragraph(text, style=style)
        self.delete_paragraph(container.paragraphs[0])

    def add_page_break(self):
        self.document.add_page_break()
