import logging
from io import BytesIO
from typing import Any

import pytest
from bs4 import BeautifulSoup
from docx import Document

from clinical_mdr_api.models import study_selections
from clinical_mdr_api.tests.integration.utils.api import inject_and_clear_db
from clinical_mdr_api.tests.integration.utils.data_library import inject_base_data
from clinical_mdr_api.tests.integration.utils.method_library import generate_study_root
from clinical_mdr_api.tests.integration.utils.utils import TestUtils
from clinical_mdr_api.tests.utils.checks import (
    assert_response_content_type,
    assert_response_status_code,
)
from clinical_mdr_api.tests.utils.utils import get_db_name
from common.config import settings

TEST_DB_NAME = get_db_name(__name__)
log = logging.getLogger(__name__)
test_data_dict: dict[str, Any]


@pytest.fixture(scope="module")
def test_database():
    log.info(
        "test_database fixture: doing ugly magic to use database: %s", TEST_DB_NAME
    )
    inject_and_clear_db(TEST_DB_NAME)
    log.info(
        "test_database fixture: injecting base data into database: %s", TEST_DB_NAME
    )
    global test_data_dict
    study, test_data_dict = inject_base_data()
    return study


# pylint: disable=unused-argument,redefined-outer-name
def test_docx_response(api_client, test_database):
    study = generate_study_root()
    response = api_client.get(f"/studies/{study.uid}/study-objectives.docx")
    assert_response_status_code(response, 200)
    assert_response_content_type(
        response,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    old_doc = Document(BytesIO(response.content))
    assert len(old_doc.tables) == 1, "DOCX document must have exactly one table"
    table = old_doc.tables[0]
    assert len(table.columns) == 4, "expected 4 columns of table"
    assert len(table.rows) == 1, "expected 1 row of table"

    # update study title to be able to lock it
    response = api_client.patch(
        f"/studies/{study.uid}",
        json={"current_metadata": {"study_description": {"study_title": "new title"}}},
    )
    assert_response_status_code(response, 200)
    # Lock
    response = api_client.post(
        f"/studies/{study.uid}/locks",
        json={
            "change_description": "Lock 1",
            "reason_for_change_uid": test_data_dict["reason_for_lock_terms"][
                0
            ].term_uid,
        },
    )
    assert_response_status_code(response, 201)
    # Unlock
    response = api_client.post(
        f"/studies/{study.uid}/unlocks",
        json={
            "reason_for_change_uid": test_data_dict["reason_for_unlock_terms"][
                0
            ].term_uid,
        },
    )
    assert_response_status_code(response, 201)

    study_objective = TestUtils.create_study_objective(
        study.uid,
        objective_template_uid=TestUtils.create_objective_template().uid,
        objective_level_uid=TestUtils.create_ct_term(
            codelist_uid=TestUtils.create_ct_codelist(
                name="Objective Level",
                sponsor_preferred_name="Objective Level",
                extensible=True,
                approve=True,
                submission_value="OBJTLEVL",
            ).codelist_uid,
            sponsor_preferred_name="level",
        ).term_uid,
    )
    unit_definition = TestUtils.create_unit_definition(name="unit1")
    timeframe = TestUtils.create_timeframe(
        timeframe_template_uid=TestUtils.create_timeframe_template().uid
    )
    TestUtils.create_template_parameter(settings.study_endpoint_tp_name)
    TestUtils.create_study_endpoint(
        study.uid,
        endpoint_template_uid=TestUtils.create_endpoint_template().uid,
        study_objective_uid=study_objective.study_objective_uid,
        endpoint_units=study_selections.study_selection.EndpointUnitsInput(
            units=[unit_definition.uid]
        ),
        timeframe_uid=timeframe.uid,
    )

    response = api_client.get(
        f"/studies/{study.uid}/study-objectives.docx?study_value_version=1"
    )
    assert_response_status_code(response, 200)
    assert_response_content_type(
        response,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    doc = Document(BytesIO(response.content))
    table = doc.tables[0]
    assert len(doc.tables) == 1, "DOCX document must have exactly one table"
    assert len(table.columns) == 4, "expected 4 columns of table"
    assert len(table.rows) == 1, "expected 1 row of table"

    response = api_client.get(f"/studies/{study.uid}/study-objectives.docx")
    assert_response_status_code(response, 200)
    assert_response_content_type(
        response,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    doc = Document(BytesIO(response.content))
    table = doc.tables[0]
    assert len(doc.tables) == 1, "DOCX document must have exactly one table"
    assert len(table.columns) == 4, "expected 4 columns of table"
    assert len(table.rows) == 3, "expected 3 rows of table"


# pylint: disable=unused-argument,redefined-outer-name
def test_html_response(api_client, test_database):
    study = generate_study_root()
    response = api_client.get(f"/studies/{study.uid}/study-objectives.html")
    assert_response_status_code(response, 200)
    assert_response_content_type(response, "text/html")
    doc = BeautifulSoup(response.content, features="lxml")
    table = doc.find("table")
    assert table, "TABLE tag not found in document"
    assert table.get("id") == "ObjectivesEndpointsTable", "TABLE id mismatch"
    assert len(table.findAll("tr")), "TABLE has no TRs"
