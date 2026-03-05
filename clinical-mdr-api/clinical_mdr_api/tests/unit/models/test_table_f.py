# pylint: disable=no-member
from typing import Mapping

import bs4
import docx
import pytest
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pyrate_limiter import Any

from clinical_mdr_api.services.utils.docx_builder import DocxBuilder
from clinical_mdr_api.services.utils.table_f import (
    Ref,
    SimpleFootnote,
    TableCell,
    TableRow,
    TableWithFootnotes,
    table_to_docx,
    table_to_html,
    table_to_xlsx,
    tables_to_html,
)

DOCX_TEXT_DIRECTION_VALUE = (
    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
)
DOCX_TEXT_DIRECTIION_TAG = (
    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}textDirection"
)

DOCX_STYLES = {
    "table": ("My Table Style", WD_STYLE_TYPE.TABLE),
    "head": ("Head Style", WD_STYLE_TYPE.PARAGRAPH),
    "data": ("Data Style", WD_STYLE_TYPE.PARAGRAPH),
}

TEST_TABLE = TableWithFootnotes(
    rows=[
        TableRow(
            hide=False,
            cells=[
                TableCell(text="hi", style="hi hi", vertical=True, footnotes=["a"]),
                TableCell(text="Hello", span="2", style="hi"),
                TableCell(style="foo", span=0),
                TableCell(text="Hello\nWorld"),
            ],
        ),
        TableRow(
            hide=True,
            cells=[
                TableCell("Blah Text", style="head2", footnotes=["x", "y"]),
                TableCell("more head", span=2),
                TableCell(span=0),
                TableCell("Also column"),
            ],
        ),
        TableRow(
            hide=False,
            cells=[
                TableCell("some Text", style="head", footnotes=["hello", "z13"]),
                TableCell("More Head", span=2, footnotes=["z"]),
                TableCell(span=0),
                TableCell("also a head"),
            ],
        ),
        TableRow(
            hide=False,
            cells=[
                TableCell("This goes TH", style="head"),
                TableCell("More TH", style="head2"),
                TableCell("data #1 cell"),
                TableCell("data cell #2", style="data"),
            ],
        ),
        TableRow(
            hide=True,
            cells=[
                TableCell("some hidden"),
                TableCell(span="2"),
                TableCell(span=0),
                TableCell("contents blah", footnotes=["hello", "a"]),
            ],
        ),
        TableRow(
            hide=False,
            cells=[
                TableCell("more Data th", vertical=True),
                TableCell(span=2),
                TableCell(span=0),
                TableCell("foo", footnotes=["z13", "hello"]),
            ],
        ),
        TableRow(
            hide=False,
            cells=[
                TableCell("foo-Bar", span=2, footnotes=["c", "X"]),
                TableCell(span=0),
                TableCell(),
                TableCell("X", footnotes=["z13", "hello"]),
            ],
        ),
        TableRow(
            hide=False,
            cells=[
                TableCell("even more\ndata"),
                TableCell("second th row"),
                TableCell("X", style="data", footnotes=["a"]),
                TableCell("data"),
            ],
        ),
    ],
    num_header_rows=3,
    num_header_cols=2,
    title="Hello Table",
    footnotes={
        "a": SimpleFootnote(uid="fn01", text_html="not used", text_plain="Footnote ej"),
        "hello": SimpleFootnote(
            uid="foot76",
            text_plain="Hello footnotes here and there",
            text_html="not displayed",
        ),
        "z13": SimpleFootnote(
            uid="foot42",
            text_plain="More footnotes after numbering",
            text_html="not shown",
        ),
    },
)

TEST_TABLE_2 = TableWithFootnotes(
    rows=[
        TableRow(
            cells=[
                TableCell(text="Visits"),
                TableCell(text="V1", refs=[Ref(type="StudyVisit", uid="visit-1")]),
                TableCell(text="V2", refs=[Ref(type="StudyVisit", uid="visit-2")]),
                TableCell(text="V3", refs=[Ref(type="StudyVisit", uid="visit-3")]),
            ]
        ),
        TableRow(
            cells=[
                TableCell(text="Activity-one"),
                TableCell(text="X", footnotes=["a"]),
                TableCell(text=""),
                TableCell(text="X"),
            ]
        ),
    ],
    num_header_rows=1,
    num_header_cols=1,
    title="Second Table",
)

EXPECTED_TABLE_HTML = """
    <table>
      <thead>
        <tr>
          <th class="hi hi">hi<sup><b>a</b></sup></th>
          <th class="hi" colspan="2">Hello</th>
          <th>Hello<br/>World</th>
        </tr>
        <tr>
          <th class="head">some Text<sup><b>hello&nbsp;z13</b></sup></th>
          <th colspan="2">More Head<sup><b>z</b></sup></th>
          <th>also a head</th>
        </tr>
      </thead>
      <tbody>
        <tr>
          <th class="head">This goes TH</th>
          <th class="head2">More TH</th>
          <td>data #1 cell</td>
          <td class="data">data cell #2</td>
        </tr>
        <tr>
          <th>more Data th</th>
          <th colspan="2"></th>
          <td>foo<sup><b>z13&nbsp;hello</b></sup></td>
        </tr>
        <tr>
          <th colspan="2">foo-Bar<sup><b>c&nbsp;X</b></sup></th>
          <td></td>
          <td>X<sup><b>z13&nbsp;hello</b></sup></td>
        </tr>
        <tr>
          <th>even more<br/>data</th>
          <th>second th row</th>
          <td class="data">X<sup><b>a</b></sup></td>
          <td>data</td>
        </tr>
      </tbody>
    </table>
    <p class="footnote"><sup><b>a</b></sup>Footnote ej</p>
    <p class="footnote"><sup><b>hello</b></sup>Hello footnotes here and there</p>
    <p class="footnote"><sup><b>z13</b></sup>More footnotes after numbering</p>
""".strip()

EXPECTED_TABLE2_HTML = """
    <table>
      <thead>
        <tr>
          <th>Visits</th>
          <th>V1</th>
          <th>V2</th>
          <th>V3</th>
        </tr>
      </thead>
      <tbody>
        <tr>
          <th>Activity-one</th>
          <td>X<sup><b>a</b></sup></td>
          <td></td>
          <td>X</td>
        </tr>
      </tbody>
    </table>
""".strip()

EXPECTED_HTML_1 = f"""<!DOCTYPE html>
<html lang="en">
  <head>
    <title>Hello Table</title>
  </head>
  <body>
    {EXPECTED_TABLE_HTML}
  </body>
</html>"""

EXPECTED_HTML_2 = f"""<!DOCTYPE html>
<html lang="en">
  <head>
    <title>Hello Table</title>
  </head>
  <body>
    {EXPECTED_TABLE_HTML}
    {EXPECTED_TABLE2_HTML}
  </body>
</html>"""


def test_table_to_html_expected():
    html = table_to_html(TEST_TABLE)
    assert html == EXPECTED_HTML_1


def test_tables_to_html_expected():
    html = tables_to_html([TEST_TABLE, TEST_TABLE_2])
    assert html == EXPECTED_HTML_2


@pytest.mark.parametrize("test_table", [TEST_TABLE])
def test_table_to_html(test_table: TableWithFootnotes):
    """Tests table_to_html by comparing table contents to TableWithFootnotes input"""

    html = table_to_html(TEST_TABLE)

    # THEN body is a parsable HTML document
    # html.parser doesn't supplement the document with additional elements for the sake of well-formedness
    doc = bs4.BeautifulSoup(html, features="html.parser")

    # THEN document has HTML tag
    assert doc.find("html"), "<HTML> tag not found"

    # THEN the document has HTML > HEAD > TITLE tags
    assert (title := doc.select_one("html > head > title")), "<TITLE> tag not found"
    # THEN document title matches
    assert title.text == test_table.title, "title mismatch"

    # THEN the document has TABLE
    assert (table := doc.select_one("html > body > table")), "<TABLE> tag not found"

    compare_html_table(table, test_table)

    if test_table.footnotes:
        compare_html_footnotes(doc, test_table)


def compare_html_table(table: bs4.element.Tag, test_table: TableWithFootnotes):
    """Compares HTML table with TableWithFootnotes by cell contents and properties"""

    # THEN number of table rows match
    expected_num_rows = sum(1 for row in test_table.rows if not row.hide)
    assert (
        len(table.find_all("tr")) == expected_num_rows
    ), "unexpected number of rows<tr>"

    # THEN number of header rows match
    num_header_rows = sum(
        1 for row in test_table.rows[: test_table.num_header_rows] if not row.hide
    )
    assert (
        len(table.select("thead > tr")) == num_header_rows
    ), "unexpected number of header rows <thead><tr>"

    row: TableRow
    # pylint: disable=invalid-name
    # iterate over index, <tr>, TableRow
    for r, (tr, row) in enumerate(
        zip(table.find_all("tr"), (row for row in test_table.rows if not row.hide))
    ):
        if r < num_header_rows:
            assert (
                tr.parent.name.lower() == "thead"
            ), f"expected parent of row {r} is <thead>"
        else:
            assert (
                tr.parent.name.lower() == "tbody"
            ), f"expected parent of row {r} is <tbody>"

        cells = [cell for cell in row.cells if cell.span]

        # THEN number of cells in rows match
        assert len(tds := tr.find_all(["th", "td"])) == len(
            cells
        ), "unexpected number of cells in row i"

        cell: TableCell
        span = 0
        # pylint: disable=invalid-name
        # iterate over index, <td/th>, TableCell
        for c, (td, cell) in enumerate(zip(tds, cells)):
            if r < num_header_rows or span < test_table.num_header_cols:
                assert (
                    td.name.lower() == "th"
                ), f"expected element in row {r} cell {c} is <th>"
            else:
                assert (
                    td.name.lower() == "td"
                ), f"expected element in row {r} cell {c} is <th>"
            span += cell.span

            # THEN cell text matches
            strings = [
                x for x in td.contents if isinstance(x, bs4.NavigableString)
            ] or [""]
            assert strings == cell.text.split(
                "\n"
            ), f"cell text doesn't match in row {r} column {c}"

            # THEN horizontal cell spanning matches
            assert td.get("colspan", "1") == str(
                cell.span
            ), f"cell span doesn't match in row {r} column {c}"

            # THEN cell styling matches
            style = " ".join(td.get("class", [])) or None
            assert (
                style == cell.style
            ), f"cell style doesn't match in row {r} column {c}"

            if cell.footnotes:
                # THEN footnote symbols found in cell as superscript
                assert (
                    len(sups := td.find_all("sup")) == 1
                ), f"expected one <sup> for footnotes in row {r} column {c}"
                symbols = sups[0].get_text(strip=True).split("\xa0")
                assert (
                    symbols == cell.footnotes
                ), f"unexpected superscript text of footnote symbols in row {r} column {c}"


# pylint: disable=invalid-name
def compare_html_footnotes(doc: bs4.BeautifulSoup, test_table: TableWithFootnotes):
    """Compares footnote <dl> with footnotes of TableWithFootnotes"""

    # THEN number of footnotes match with footnotes list
    assert (paras := doc.find_all("p", class_="footnote")), "no P.footnote tags found"
    assert len(paras) == len(test_table.footnotes), "number of footnotes mismatch"

    for r, (para, (symbol, footnote)) in enumerate(
        zip(paras, test_table.footnotes.items())
    ):
        # THEN footnote symbol is super bold
        assert (supr := para.find("sup")), f"footnote symbol is not super in row {r}"
        assert (bold := supr.find("b")), f"footnote symbol is not bold in row {r}"

        # THEN footnote symbols matches
        assert (
            bold.get_text() == symbol
        ), f"<dt> text doesn't match symbol in footnote row {r}"

        # THEN footnote text matches
        assert (
            para.get_text() == f"{symbol}{footnote.text_plain}"
        ), f"footnote plain-text doesn't match in row {r}"


@pytest.mark.parametrize("test_table", [TEST_TABLE])
def test_table_to_docx(test_table: TableWithFootnotes):
    """Tests table_to_docx() by comparing DOCX document to TableWithFootnotes input"""

    docx_doc: DocxBuilder = table_to_docx(test_table, styles=DOCX_STYLES).document

    # THEN the document contains exactly one table
    assert len(docx_doc.tables) == 1, "expected exactly 1 table in DOCX SoA"

    compare_docx_table(docx_doc.tables[0], test_table, DOCX_STYLES)

    if test_table.footnotes:
        compare_docx_footnotes(docx_doc, test_table.footnotes, DOCX_STYLES)


def compare_docx_table(
    tablex: docx.table.Table,
    test_table: TableWithFootnotes,
    docx_styles: Mapping[str, tuple[str, Any]],
):
    """Compares DOCX table with TableWithFootnotes by column contents and properties"""

    # THEN number of table rows match
    expected_num_rows = sum(1 for row in test_table.rows if not row.hide)
    assert len(tablex.rows) == expected_num_rows, "unexpected number of table rows"

    # THEN number of table columns match
    expected_num_cols = sum(cell.span for cell in test_table.rows[-1].cells)
    assert (
        len(tablex.columns) == expected_num_cols
    ), "unexpected number of table columns"

    row: TableRow
    rowx: docx.table._Row
    for row_idx, (rowx, row) in enumerate(
        zip(tablex.rows, (row for row in test_table.rows if not row.hide))
    ):
        cell: TableCell
        cellx: docx.table._Cell
        for col_idx, (cellx, cell) in enumerate(zip(rowx.cells, row.cells)):
            if not cell.span:
                # In DocX, despite merging cells, the number of cells in a row is constant across the table, and the
                # contents of the first merged member are repeated for all members.
                # So when we encounter a span=0 cell in JSON, it's a subsequent member of a merged cell,
                # we skip checking the contents and formatting of the corresponding cell in DocX.
                continue

            # first paragraph of cell
            parax0: docx.text.paragraph.Paragraph = cellx.paragraphs[0]

            # THEN cell text matches
            textx = parax0.runs[0].text if parax0.runs else parax0.text
            assert (
                textx == cell.text
            ), f"cell text doesn't match in row {row_idx} column {col_idx}"

            if cell.vertical:
                # THEN some header cell texts are oriented vertical
                tcpr = next(
                    (
                        tcPr
                        for tcPr in cellx._tc.tcPr
                        if tcPr.tag == DOCX_TEXT_DIRECTIION_TAG
                    ),
                    None,
                )
                assert (
                    tcpr is not None
                ), f"vertical cell direction not set on row {row_idx} column {col_idx}"
                assert (
                    tcpr.get(DOCX_TEXT_DIRECTION_VALUE) == "btLr"
                ), f"vertical cell direction 'btLr' expected row {row_idx} column {col_idx}"

            # THEN header columns are aligned left, other columns centered
            if col_idx < test_table.num_header_cols:
                assert (
                    parax0.alignment is None
                ), f"paragraph is aligned in row {row_idx} column {col_idx}"
            else:
                assert (
                    parax0.alignment == WD_ALIGN_PARAGRAPH.CENTER
                ), f"paragraph is not centered in row {row_idx} column {col_idx}"

            if cell.style in docx_styles:
                # THEN cell styling matches
                expected_style_name = docx_styles.get(cell.style, ["Normal"])[0]
                assert (
                    parax0.style.name == expected_style_name
                ), f"cell style doesn't match in row {row_idx} column {col_idx}"

            if cell.footnotes:
                # THEN footnote symbols match
                symbols = parax0.runs[1].text.split("\u00A0")
                assert (
                    symbols == cell.footnotes
                ), f"footnote symbols don't match in {row_idx} column {col_idx}"

                # THEN footnote symbols are in bold superscript
                assert all(
                    runx.font.superscript for runx in parax0.runs[1:]
                ), f"footnote symbols are not superscript in row {row_idx} column {col_idx}"
                assert all(
                    runx.font.bold for runx in parax0.runs[1:]
                ), f"footnote symbols are not bold in row {row_idx} column {col_idx}"


def compare_docx_footnotes(
    docx_doc: DocxBuilder,
    footnotes: dict[str, SimpleFootnote],
    docx_styles: Mapping[str, tuple[str, Any]],
):
    """Compares DOCX paragraphs with footnotes of TableWithFootnotes"""

    # Table is not a paragraph, so all top-level paragraphs of a SoA DOCX document are part of footnote listing
    footnote_paras = [parax for parax in docx_doc.paragraphs if parax.text.strip()]

    # THEN document lists footnotes
    assert len(footnote_paras), "footnotes listing was not found in DOCX document"

    # THEN number of footnotes match
    assert len(footnote_paras) == len(footnotes), "number of footnotes mismatch"

    for row_idx, (parax, (symbol, footnote)) in enumerate(
        zip(footnote_paras, footnotes.items())
    ):
        # THEN footnote paragraph has styling
        assert parax.style.name == docx_styles.get("footnote", ["Normal"])[0]

        # THEN footnote symbols match with footnotes list
        textx = parax.runs[0].text
        assert textx == symbol, f"footnote symbol doesn't match in row {row_idx}"

        # THEN footnote symbols are bold superscript
        assert parax.runs[
            0
        ].font.superscript, f"footnote symbol is not superscript in row {row_idx}"
        assert parax.runs[0].font.bold, f"footnote symbol is not bold in row {row_idx}"

        # Then footnote text matches with footnotes list
        textx = parax.runs[1].text
        assert (
            textx == footnote.text_plain
        ), f"footnote text doesn't match in row {row_idx}"


def test_table_to_xlsx():
    workbook = table_to_xlsx(TEST_TABLE)
    worksheet = workbook.active

    assert worksheet.title == TEST_TABLE.title
    assert worksheet.max_row == len(TEST_TABLE.rows)
    assert worksheet.max_column == len(TEST_TABLE.rows[0].cells)
    assert worksheet.freeze_panes == "C4"

    for r, row in enumerate(TEST_TABLE.rows, start=1):
        for c, cell in enumerate(row.cells, start=1):
            value = worksheet.cell(row=r, column=c).value
            if value is None:
                value = ""
            assert value == cell.text

    expected_merged = {
        "B1:C1",
        "B2:C2",
        "B3:C3",
        "B5:C5",
        "B6:C6",
        "A7:B7",
    }
    merged_ranges = {str(rng) for rng in worksheet.merged_cells.ranges}
    assert expected_merged.issubset(merged_ranges)
